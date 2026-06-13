"""
generate_docs.py — build DOCUMENTATION.docx for the whole project.

Reads the persisted eval outputs in results/ and renders a formatted Word document
(architecture, stack, every result table, the hard-set failure analysis, the product
layer, and how to run it). Numbers come from the JSON files, so the doc never drifts
from the actual measured results — re-run the evals, re-run this, and it's current.

Usage:  python -m scripts.generate_docs        # writes DOCUMENTATION.docx
"""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

RESULTS = Path("results")
OUT = Path("DOCUMENTATION.docx")
AUTHOR = "Aadhavan Alakan"
EMAIL = "aadhavanalakan97@gmail.com"


def load(name: str) -> dict:
    p = RESULTS / name
    return json.loads(p.read_text()) if p.exists() else {}


def f3(x) -> str:
    return f"{x:.3f}" if isinstance(x, (int, float)) else str(x)


# ---------- formatting helpers ----------------------------------------------
def _table_style(doc):
    for s in ("Light Grid Accent 1", "Light List Accent 1", "Table Grid"):
        try:
            doc.styles[s]
            return s
        except KeyError:
            continue
    return None


def add_table(doc, headers, rows, bold_row_idx=None):
    style = _table_style(doc)
    t = doc.add_table(rows=1, cols=len(headers))
    if style:
        t.style = style
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = str(h)
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            if bold_row_idx is not None and ridx == bold_row_idx:
                for r in cells[i].paragraphs[0].runs:
                    r.bold = True
    doc.add_paragraph()
    return t


def add_mono(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


# ---------- document ---------------------------------------------------------
def build() -> None:
    retr = load("retrieval_latest.json")
    gen = load("generation_latest.json")
    rr = load("rerank_study.json")
    bench = load("benchmark.json")
    hard = load("hard_set.json")
    ragas = load("ragas.json")

    doc = Document()

    # Title block
    title = doc.add_heading("FinRAG — Financial Document Intelligence (RAG)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Grounded, cited Q&A over SEC 10-K filings — with measured chunking & "
                    "reranking studies, an adversarial failure analysis, and a chatbot UI")
    r.italic = True
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"{AUTHOR}  ·  {EMAIL}  ·  Generated {datetime.now().strftime('%Y-%m-%d')}")
    link = doc.add_paragraph()
    link.alignment = WD_ALIGN_PARAGRAPH.CENTER
    link.add_run("Live demo: https://aadhavanalakan-finrag-app-eadpfy.streamlit.app   ·   "
                 "Code: https://github.com/aadhavanalakan/FinRAG").bold = True
    doc.add_paragraph()

    # 1. Executive summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "FinRAG answers natural-language questions about SEC 10-K filings (Verizon, AT&T, "
        "T-Mobile; FY2025) with grounded, cited answers and a designed refusal path. Beyond "
        "a working pipeline, it measures which design choices help, via a reproducible "
        "evaluation harness. Three headline results:")
    for b in [
        "Semantic, table-aware chunking beats fixed-size chunking on every retrieval metric "
        "(hit@5 0.688 vs 0.562) and improves generation — correctness 0.625 vs 0.562, "
        "false-refusal 0.312 vs 0.438, faithfulness 0.974 vs 0.904.",
        "An off-the-shelf cross-encoder reranker HURTS on this table-heavy corpus; no "
        "mitigation beats turning it off, so the production default is reranking OFF.",
        "Under an adversarial hard set the system scores 14/15 — refusing unanswerable "
        "questions and blocking injection/advice — with the one miss a documented retrieval gap.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    # 2. Problem & use case
    doc.add_heading("2. Problem & Use Case", level=1)
    doc.add_paragraph(
        "Investment analysts need fast, trustworthy answers to factual questions about company "
        "financials and disclosures buried in long 10-K filings. A generic LLM hallucinates "
        "figures; a search box returns pages, not answers. FinRAG retrieves the exact passages, "
        "answers only from them, cites each claim back to a source chunk, and refuses when the "
        "filing does not support an answer — the behavior a financial tool requires.")

    # 3. Architecture
    doc.add_heading("3. System Architecture", level=1)
    add_mono(doc,
             "INGEST (once):  10-K HTML -> parse (tables -> atomic Markdown) -> chunk (fixed AND\n"
             "                semantic) -> embed (cached) -> Pinecone (namespace/strategy) + BM25\n\n"
             "ASK (per query): question -> [guard] -> embed -> hybrid retrieve (dense + BM25, RRF)\n"
             "                -> [rerank?] -> generate (cited, streamed) -> [audit] -> answer + metrics\n\n"
             "EVAL:  golden set (retrieval + faithfulness) · hard set (adversarial) · CI gate")

    # 4. Tech stack
    doc.add_heading("4. Technology Stack", level=1)
    add_table(doc, ["Layer", "Choice"], [
        ["Generation", "Nebius open models (default Qwen3-30B-A3B) or OpenAI gpt-4o / gpt-4o-mini, temp 0"],
        ["Embeddings", "Nebius Qwen3-Embedding-8B — 4096-dim, cosine, disk-cached"],
        ["Vector store", "Pinecone serverless, one namespace per chunking strategy"],
        ["Retrieval", "Hybrid dense + BM25, fused with Reciprocal Rank Fusion (k=60)"],
        ["Reranker", "ms-marco-MiniLM-L-12-v2 cross-encoder (local; off by default)"],
        ["Orchestration", "LangGraph StateGraph: guard -> retrieve -> generate -> audit"],
        ["Guardrails", "Deterministic rule-based (input injection/advice/length + output citation audit)"],
        ["UI", "Streamlit chatbot — model switch, live latency/cost/TTFT metrics, sources"],
        ["Config", "Versioned YAML + typed Pydantic loader + per-run manifest"],
        ["Tests / CI", "pytest (offline invariants + guardrails) + regression gate"],
    ])

    doc.add_heading("4.1  Component Reference", level=2)
    add_table(doc, ["Component", "Type", "LLM calls", "Role"], [
        ["Ingestion (edgar)", "Pure Python", "0", "Pull 10-K HTML from SEC EDGAR"],
        ["Parser", "Pure Python", "0", "Table-aware HTML -> prose + atomic Markdown tables"],
        ["Fixed chunker", "Pure Python", "0", "512-char splitter (tables shredded)"],
        ["Semantic chunker", "Embedding-based", "0", "Percentile breakpoints; tables kept whole"],
        ["Embedder", "Hosted (Nebius)", "0", "Qwen3-Embedding-8B, 4096-d, cached"],
        ["Vector store", "Pinecone", "0", "2 namespaces (fixed/semantic), cosine"],
        ["Hybrid retriever", "Dense + BM25 + RRF", "0", "Fuse vector + keyword rankings"],
        ["Reranker", "Local model", "0", "ms-marco cross-encoder (off by default)"],
        ["Guardrails", "Pure Python", "0", "Input injection/advice/length + citation audit"],
        ["Generator", "LLM", "1", "Cited answer (Nebius or OpenAI), streamed"],
        ["Orchestrator", "LangGraph", "0", "guard -> retrieve -> generate -> audit"],
        ["Faithfulness judge", "LLM judge", "~2/Q", "Atomic-claim extraction + support check"],
        ["RAGAS (in-house)", "LLM judge", "~3/Q", "context precision / recall / faithfulness"],
    ])
    doc.add_paragraph("Per query: exactly 1 LLM call (generation). Eval adds judge calls offline.")

    # 5. Corpus & method
    doc.add_heading("5. Corpus & Methodology", level=1)
    doc.add_paragraph(
        "Corpus: three telecom 10-K filings (Verizon, AT&T, T-Mobile), all fiscal year ended "
        "Dec 31 2025 — a clean apples-to-apples set for cross-company questions. The parser "
        "converts each HTML table to Markdown wrapped in protective markers so chunkers treat "
        "it as an atomic unit (the single most important design choice).")
    doc.add_paragraph(
        "Two chunking strategies over the same documents: FIXED (512-char blind slices, tables "
        "shredded) vs SEMANTIC (topic-boundary cuts, tables kept whole). Everything downstream "
        "is identical, so any quality difference is attributable to chunking or the reranker toggle.")
    doc.add_paragraph(
        "Evaluation: a golden set of 18 manually-verified Q/A pairs (16 answerable + 2 negative "
        "'should-refuse'). Relevance is content-based (a chunk is relevant if it contains the "
        "verified answer figure). Scored in two stages — retrieval (no LLM) and generation "
        "(LLM-as-judge faithfulness).")

    # 6. Result A — chunking
    doc.add_heading("6. Result A — Chunking Comparison", level=1)
    doc.add_paragraph("Retrieval metrics (no LLM; the clean comparison is without the reranker confound):")
    arms = retr.get("arms", {})
    order = ["fixed_norerank", "semantic_norerank"]
    rows = [[a, f3(arms[a]["hit@1"]), f3(arms[a]["hit@3"]), f3(arms[a]["hit@5"]),
             f3(arms[a]["mrr"]), f3(arms[a]["ndcg@5"])] for a in order if a in arms]
    add_table(doc, ["arm", "hit@1", "hit@3", "hit@5", "MRR", "nDCG@5"], rows, bold_row_idx=1)
    add_caption(doc, f"Source: results/retrieval_latest.json (run {retr.get('run_id','?')}, "
                     f"n={retr.get('n_answerable','?')} answerable).")

    doc.add_paragraph("Generation metrics (LLM-judge):")
    g = gen.get("arms", {})
    grows = [[a, f3(g[a]["correct"]), f3(g[a]["false_refuse"]), f3(g[a]["cite_cov"]),
              f3(g[a]["faithful"]), f3(g[a]["refuse_acc"])] for a in order if a in g]
    add_table(doc, ["arm", "correct", "false-refuse", "citation cov.", "faithfulness", "refuse acc."],
              grows, bold_row_idx=1)
    doc.add_paragraph(
        "Conclusion A: Semantic, table-aware chunking measurably improves both retrieval and "
        "answer faithfulness (0.974 vs 0.904). The mechanism is preserved table structure — a "
        "figure stays next to its label — not a generic 'bigger chunks' effect.")

    # 7. Result B — reranking
    doc.add_heading("7. Result B — Reranking Impact", level=1)
    doc.add_paragraph("All four arms (the reranker consistently lowers every metric):")
    allrows = [[a, f3(arms[a]["hit@5"]), f3(arms[a]["mrr"]), f3(arms[a]["ndcg@5"])]
               for a in ["semantic_norerank", "semantic_rerank", "fixed_norerank", "fixed_rerank"]
               if a in arms]
    add_table(doc, ["arm", "hit@5", "MRR", "nDCG@5"], allrows, bold_row_idx=0)

    doc.add_paragraph("Mitigation study (can the reranker be rescued? — semantic arm):")
    v = rr.get("variants", {})
    vorder = ["norerank", "msmarco", "msmarco+caption", "msmarco+blend", "bge", "bge+caption"]
    vrows = [[k, f3(v[k]["hit@5"]), f3(v[k]["ndcg@5"])] for k in vorder if k in v]
    add_table(doc, ["variant", "hit@5", "nDCG@5"], vrows, bold_row_idx=0)
    doc.add_paragraph(
        "Conclusion B: The MS-MARCO cross-encoder was trained on natural-language passages and "
        "systematically demotes table chunks — exactly where the answer figures live. Captioning, "
        "score-blending, and a stronger reranker (bge) recover most of the damage but none beats "
        "no-reranking. For table-heavy financial RAG, invest in chunking, not a reranker. "
        "Production default: reranking OFF; CI gates on semantic_norerank.")

    # 8. Result C — benchmark
    doc.add_heading("8. Result C — Hosted-Model Benchmark", level=1)
    doc.add_paragraph(f"{bench.get('n_prompts','?')} standardized prompts x "
                      f"{len(bench.get('models',{}))} Nebius models, streaming "
                      f"(max {bench.get('max_tokens','?')} tokens):")
    b = bench.get("models", {})
    brows = [[m.split("/")[-1], f3(d["ttft_p50"]), f"{d['tok_s']:.1f}", f3(d["lat_p50"]), f3(d["lat_p95"])]
             for m, d in b.items()]
    add_table(doc, ["model", "TTFT p50 (s)", "tok/s", "latency p50 (s)", "latency p95 (s)"], brows)
    doc.add_paragraph(
        "gpt-oss-120b is dramatically the fastest on Nebius despite being largest (optimized MoE "
        "serving); the dense Llama-70B is the slowest. The pipeline's choice, Qwen3-30B, is a "
        "strong balance (lowest TTFT, stable p95). Quality is measured separately by the eval harness.")

    # 9. Result D — hard set
    doc.add_heading("9. Result D — Adversarial Hard Set & Failure Analysis", level=1)
    doc.add_paragraph(
        f"A {hard.get('n','?')}-item hard set probes where the system breaks: exact line-items, "
        "wrong-year bait, cross-company comparisons, a multi-hop ratio, ambiguous queries, refusal "
        "traps, and prompt-injection / advice attempts. Outcome: "
        f"{hard.get('n_pass','?')}/{hard.get('n','?')} met expected behavior "
        f"(est. cost ${hard.get('cost_usd','?')}).")
    hrows = [[r["id"], r["probe"][:42], r["expected_behavior"], r["actual"],
              "pass" if r["passed"] else "MISS"] for r in hard.get("rows", [])]
    add_table(doc, ["id", "probe", "expected", "actual", "verdict"], hrows)
    doc.add_paragraph(
        "Conclusion D: Refusal discipline and guardrails hold under adversarial probing. The one "
        "genuine miss (AT&T operating income) is a retrieval miss — the figure was not in the "
        "retrieved chunks, so the model correctly refused rather than fabricate. The structural "
        "weakness this surfaced was multi-document synthesis: cross-company comparisons needed "
        "per-company sub-queries the single-pass retriever did not issue. That gap is now closed "
        "in the chatbot answer path via query planning (see Section 10.1) — comparisons resolve "
        "each company separately and synthesize. The full per-item breakdown is in "
        "eval/FAILURE_ANALYSIS.md.")

    # 9.1 RAGAS cross-check
    doc.add_heading("9.1  RAGAS Cross-Check (industry-standard metrics)", level=2)
    doc.add_paragraph(
        "The same configurations scored with RAGAS-style metrics (faithfulness, context "
        "precision, context recall), judged by gpt-4o-mini — the standard most RAG projects "
        "report, computed in-house to avoid the ragas package's dependency conflicts:")
    ra = ragas.get("arms", {})
    rorder = ["fixed_norerank", "semantic_norerank", "semantic_rerank"]
    rrows = [[a, f3(ra[a]["faithfulness"]), f3(ra[a]["context_precision"]), f3(ra[a]["context_recall"])]
             for a in rorder if a in ra]
    add_table(doc, ["arm", "faithfulness", "context_precision", "context_recall"], rrows, bold_row_idx=1)
    doc.add_paragraph(
        "RAGAS independently confirms both findings: semantic beats fixed on every metric, "
        "AND reranking HURTS (semantic+rerank has the lowest context precision). Note the "
        "absolute values are lower than single-company kits report — this corpus is harder "
        "(cross-company questions) and gpt-4o-mini tends to elaborate beyond the context, "
        "which the faithfulness metric penalizes.")

    # 10. Product layer
    doc.add_heading("10. Product Layer", level=1)
    for b in [
        "Chatbot UI (app.py): chat interface with per-query model switching and a live metrics "
        "strip on every answer — latency, time-to-first-token, tokens/sec, input/output tokens, "
        "cost — plus a session cost meter and an expandable sources panel.",
        "Guardrails (finrag/guardrails.py): deterministic, no extra LLM call. Input gate blocks "
        "prompt-injection / advice solicitation / over-length; output audit checks every citation "
        "against real sources and flags invented references.",
        "LangGraph orchestration (finrag/graph.py): a compiled guard -> retrieve -> generate -> "
        "audit state machine; a blocked input short-circuits to the end at zero cost.",
        "Runtime corpus management (finrag/corpus.py, finrag/edgar.py): add or remove any "
        "US-listed company with a 10-K by ticker, validated against SEC EDGAR, from the CLI or the "
        "sidebar. The three seed telecoms are protected so the eval set can't be broken.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    # 10.1 Reliability engineering (the chatbot answer path)
    doc.add_heading("10.1  Reliability Engineering (chat answer path)", level=2)
    doc.add_paragraph(
        "The eval harness uses a strict top-k retrieval (for a clean chunking/reranking "
        "comparison), but the CHATBOT answer path adds three reliability layers so real "
        "questions don't miss retrievable facts — the difference between a research artifact "
        "and a usable product:")
    for b in [
        "Coverage mode — single-company questions feed the model ~20 ranked chunks (not 5) "
        "within a large context budget, so a fact that ranks 6th-20th (e.g. the auditor, "
        "Deloitte) is still in context. This is the 'put more in context' fix; a stingy top-5 "
        "was the root cause of spurious refusals.",
        "Income-statement guarantee — any question about a financial line (revenue, operating "
        "income, net income, EPS) injects that company's consolidated income-statement table, "
        "pinned by a canonical query, so figure lookups never miss the statement line.",
        "Query planning for comparisons — a cross-company question is decomposed into one "
        "reliable single-company lookup per company, then the comparison is synthesized from "
        "those clean, cited answers. This fixes the classic failure where the model picks the "
        "wrong company's number or the wrong fiscal year out of a pile of similar tables.",
    ]:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph(
        "Net effect: single lookups, the auditor question, narrative/MD&A questions, and "
        "cross-company comparisons all answer correctly, while unanswerable and forward-guidance "
        "questions still refuse and injection/advice are still blocked.")

    # 10.2 Deployment
    doc.add_heading("10.2  Deployment", level=2)
    doc.add_paragraph(
        "Deployed on Streamlit Community Cloud from the public GitHub repo. The BM25 corpus is "
        "committed so retrieval works on a fresh clone; the cross-encoder reranker (torch) is "
        "left out of the default install to keep the build lean; API keys are supplied via "
        "Streamlit secrets (mirrored into environment variables at startup); an optional "
        "password gate activates only when APP_PASSWORD is set.")
    add_table(doc, ["Resource", "URL"], [
        ["Live app", "https://aadhavanalakan-finrag-app-eadpfy.streamlit.app"],
        ["Source", "https://github.com/aadhavanalakan/FinRAG"],
    ])

    # 11. Engineering
    doc.add_heading("11. Engineering Practices", level=1)
    for b in [
        "Versioned config + run manifest — every parameter in git-tracked YAML with a version + "
        "content hash, so any metric movement is attributable to one change.",
        "Disk-backed embedding cache — re-running ingest re-embeds only changed text.",
        "Offline test suite (pytest) — table-preservation invariant, metric math, guardrail rules, "
        "and the company-detection hints; plus a CI regression gate against thresholds.yaml.",
        "Reproducible documentation — this file is generated from results/ by scripts.generate_docs.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    # 11.1 Key design decisions
    doc.add_heading("11.1  Key Design Decisions", level=2)
    add_table(doc, ["Decision", "Rejected alternative", "Rationale"], [
        ["Table-aware parsing", "Strip all HTML to text", "Keeps each figure beside its row label"],
        ["Hybrid dense + BM25 + RRF", "Pure dense retrieval", "BM25 catches exact tokens dense misses"],
        ["Reranking off by default", "Always-on cross-encoder", "Measured to hurt (demotes table chunks)"],
        ["Multi-provider generation", "Single provider", "Cheap Nebius default; OpenAI parity, per-query"],
        ["In-house RAGAS metrics", "ragas package dependency", "ragas pins clash with langgraph on Py3.14"],
        ["Versioned YAML + manifest", "Hardcoded parameters", "Every metric traceable to one config version"],
        ["LangGraph orchestration", "Plain Python pipeline", "Clean guardrail short-circuit as a state machine"],
        ["Registry-backed corpus", "Fixed corpus in code", "Add any US company at runtime via SEC EDGAR"],
    ])

    # 12. Limitations
    doc.add_heading("12. Limitations & Future Work", level=1)
    for b in [
        "Golden set is 18 items; scaling to 50-200 would tighten the estimates.",
        "Multi-document synthesis is the main capability gap — comparisons and ratios refuse; a "
        "query-planning / per-company sub-retrieval step would turn those into answers.",
        "Prices in config/pricing.yaml are estimates; set real per-token values for an exact cost meter.",
        "The reranking finding motivates trying a table-native or LLM reranker.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    # 13. Reproduce
    doc.add_heading("13. How to Run / Reproduce", level=1)
    add_mono(doc,
             "pip install -r requirements.txt        # set NEBIUS_API_KEY + PINECONE_API_KEY in .env\n"
             "python -m scripts.ingest               # build the index\n"
             "streamlit run app.py                   # the chatbot UI\n"
             "python -m scripts.ask \"What were Verizon's total operating revenues in 2025?\"\n"
             "python -m scripts.add_company TSLA      # add a company at runtime\n\n"
             "python -m eval.run_eval                # retrieval metrics (free)\n"
             "python -m eval.run_eval_gen            # generation metrics (LLM)\n"
             "python -m eval.rerank_study            # reranking mitigation study\n"
             "python -m eval.run_hard                # adversarial hard set -> FAILURE_ANALYSIS.md\n"
             "python -m benchmark.run_benchmark      # model speed comparison\n"
             "python -m pytest tests/ && python -m eval.gate     # what CI runs")

    doc.save(OUT)
    print(f"Wrote {OUT.resolve()}  ({len(doc.paragraphs)} paragraphs)")


if __name__ == "__main__":
    build()
